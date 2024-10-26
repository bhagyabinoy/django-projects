from django.shortcuts import render
from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from .serializers import *
from .models  import *
from openpyxl import load_workbook
from rest_framework.parsers import JSONParser 
import pdfkit
from django.http import HttpResponse
import logging
import datetime
from django.template.loader import render_to_string
from .utils import food_utils, log_utils
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Create your views here.
##################################################################################################################################
logger = logging.getLogger(__name__)
#RETRIEVE ALL category
class CategoryList(APIView):

    permission_classes = [AllowAny]
    
    def get(self, request): 
            #queryset = CategoryModel.objects.all().order_by('id')
            queryset = food_utils.get_all_category_instances()
            serializer = CategorySerializer(queryset, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
   

##################################################################################################################################

#REGISTERcategory    
class CategoryCreate(APIView):

    permission_classes = [IsAuthenticated]

    def post(self, request): 

        if(request.user.is_staff):  

            serializer = CategorySerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                category=serializer.data,
                return Response(category, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)    

##################################################################################################################################

#DETAIL category  
class CategoryDetail(APIView):
    
    #permission_classes = [IsAuthenticated]
    permission_classes = ([AllowAny])
    def get(self,request, pk):

        if request.user.is_staff:

            #queryset = CategoryModel.objects.get(id=pk)
            queryset = food_utils.category_detail(pk)
            serializer = CategorySerializer(queryset, many=False)
            return Response(serializer.data, status=status.HTTP_200_OK)
        
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)   
        
##################################################################################################################################

#UPDATE category

class CategoryUpdate(APIView):

    permission_classes = [AllowAny]
    #permission_classes = [IsAuthenticated]

    def patch(self,request, pk):

        if request.user.is_staff:

            #queryset = CategoryModel.objects.get(id=pk)
            queryset = food_utils.category_detail(pk)
            serializer = CategorySerializer(instance=queryset, data=request.data,partial=True)
            if serializer.is_valid():

                serializer.save()
                return Response(serializer.data, status=status.HTTP_200_OK)
            
            else:
                return Response(serializer.errors,status=status.HTTP_304_NOT_MODIFIED)
            
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)   

##################################################################################################################################

#DELETE category

class CategoryDelete(APIView):

    permission_classes = [AllowAny]
    #permission_classes = ([AllowAny])
    def get(self, request, pk):

        if request.user.is_staff:
            #queryset = CategoryModel.objects.get(id=pk)
            queryset = food_utils.category_detail(pk)
            queryset.delete()
            return Response("Category deleted successfully", status=status.HTTP_204_NO_CONTENT)
        
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)   
    
##################################################################################################################################

#RETRIEVE ALL USERS
class FoodList(APIView):

    permission_classes = [IsAuthenticated]
    
    def get(self, request): 

            #queryset = FoodItem.objects.all().order_by('id')
            queryset = food_utils.get_all_food_item_instances()
            serializer = FoodSerializer(queryset, many=True)
            logger.warning('Food page was accessed at '+str(datetime.datetime.now())+' hours!')
            return Response(serializer.data, status=status.HTTP_200_OK)
   

##################################################################################################################################

#REGISTER USER    
class FoodCreate(APIView):

    permission_classes = [IsAuthenticated]
  
    def post(self, request):     

        if request.user.is_staff:  
            serializer = FoodSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response({"message": "Food Created successfully", "food": serializer.data},status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)    

##################################################################################################################################

#DETAIL USER   
class FoodDetail(APIView):
    
    permission_classes = [IsAuthenticated]

    def get(self,request, pk):

        if request.user.is_staff:
            try:
                #queryset = FoodItem.objects.get(id=pk)
                queryset = food_utils.food_detail(pk)

            except FoodItem.DoesNotExist:
                return Response("No Food item was found for the given ID", status=status.HTTP_404_NOT_FOUND)
            serializer = FoodSerializer(queryset, many=False)
            print(serializer.data)         # {'id': 21, 'category': 'NonVeg', 'name': 'cheese', 'image': '/media/images/images/download_8.jpeg', 'quantity': 100, 'calories': 402}
            return Response(serializer.data, status=status.HTTP_200_OK)
        
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)   
        
##################################################################################################################################

#UPDATE USER

class FoodUpdate(APIView):

    permission_classes = [IsAuthenticated]

    def patch(self,request, pk):

        if request.user.is_staff:
            try:
                #queryset = FoodItem.objects.get(id=pk)
                queryset = food_utils.food_detail(pk)
            except FoodItem.DoesNotExist:
                return Response("No Food item was found for the given ID", status=status.HTTP_404_NOT_FOUND)
            
            serializer = FoodSerializer(instance=queryset, data=request.data,partial=True)
            print(request.data)
            if serializer.is_valid(raise_exception=ValueError):
                serializer.save()
                return Response(serializer.data, status=status.HTTP_200_OK)
            else:
                return Response(serializer.errors,status=status.HTTP_304_NOT_MODIFIED)
            
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)   

##################################################################################################################################

#DELETE USER

class FoodDelete(APIView):

    permission_classes = [IsAuthenticated]
    #permission_classes = ([AllowAny])
    def get(self, request, pk):
        
        if request.user.is_staff:
            try:
                #queryset = FoodItem.objects.get(id=pk)
                queryset = food_utils.food_detail(pk)

            except FoodItem.DoesNotExist:
                return Response("No Food item was found for the given ID", status=status.HTTP_404_NOT_FOUND)
            
            queryset.delete()
            return Response("Food Item deleted successfully", status=status.HTTP_204_NO_CONTENT)
        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)   
    
##################################################################################################################################
"""FILTER"""

#BASED ON FOOD CATEGORY

class CategoryFilter(APIView):

    permission_classes = [AllowAny]

    def get(self, request): 
        category = request.query_params.get("category")
        print(category)
        if category=='Veg':
            #queryset = FoodItem.objects.filter(category_id=1)
            queryset = food_utils.filter_veg()
            print(queryset)
            serializer = FoodSerializer(queryset, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
        elif category=='NonVeg':
            #queryset = FoodItem.objects.filter(category_id=2)
            queryset = food_utils.filter_nonveg()
            serializer = FoodSerializer(queryset, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
        else:
            return Response("No Food item in this category is available", status=status.HTTP_404_NOT_FOUND)   


##################################################################################################################################

#LESS THAN 100

class Caloryless(APIView):

    permission_classes = [AllowAny]

    def get(self, request): 

        #queryset = FoodItem.objects.all().filter(calories__lt=100)
        queryset = food_utils.calories_lt_100()
        print(queryset)        
        if queryset.count() == 0:		
                return Response("No Food item was found", status=status.HTTP_404_NOT_FOUND)     
        else: 
            serializer = FoodSerializer(queryset, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
                

##################################################################################################################################

#B/W 100-200
class CaloryGt(APIView):

    permission_classes = [AllowAny]

    def get(self, request): 
        #queryset = FoodItem.objects.all().filter(calories__gt=300)
        queryset = FoodItem.objects.all().filter(calories__gt=300)

        print(queryset)        
        if queryset.count() == 0:		
            return Response("No Food item was found", status=status.HTTP_404_NOT_FOUND)     
        else: 
            serializer = FoodSerializer(queryset, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
    

##################################################################################################################################

#GREATER THAN 200
class CaloryRange(APIView):

    permission_classes = [AllowAny]

    def get(self, request): 
        #queryset = FoodItem.objects.all().filter(calories__range=(100, 200))
        queryset =  food_utils.calories_range()
        if queryset.count() == 0:		
            return Response("No Food item was found", status=status.HTTP_404_NOT_FOUND)     
        else: 
            serializer = FoodSerializer(queryset, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)
    

##################################################################################################################################

#SEARCH
class SearchFood(APIView):

    permission_classes = [AllowAny]

    def get(self, request): 
        foodname = request.query_params.get("foodname")
        queryset = FoodItem.objects.filter(name__icontains=foodname).values()
        serializer = FoodsearchSerializer(queryset, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    

##################################################################################################################################

#BULK CREATE
     
class Bulkcreate(APIView):

    def post(self, request):

        if request.user.is_staff:

            file_name = request.FILES['file']  

            wb = load_workbook(file_name)
            sheet = wb.active

            key = []  
            for col in range(1, sheet.max_column+1):
                key.append(sheet.cell(row=1, column=col).value)
            data =[]
            for row in range(2, sheet.max_row+1):
                dict = {}
                for col in range(1, sheet.max_column+1):
                    dict[key[col-1]]=sheet.cell(row=row,column=col).value
                data.append(dict)
                #print(dict)

            #header checking
            header = ['name', 'category', 'quantity', 'calories']
            if len(key) != len(header):
                return Response({'message':'Fields are not identical'}, status=status.HTTP_400_BAD_REQUEST)
            else:
                List = []
                for i in header:
                    if i in key:
                        List.append(i)
                        if len(key) != len(header):
                            return Response({'message':'Invalid Field(s) present in file'}, status=status.HTTP_400_BAD_REQUEST)
    
            bulk_request = data
            print(bulk_request)
            wb.close()

            serializer =  FoodSerializer(data=bulk_request, many=True)
            if serializer.is_valid(raise_exception=ValueError):
                serializer.save()
                print(serializer.data)
                return Response(serializer.data,  status=status.HTTP_201_CREATED)
            else:
                print(serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        else:
            return Response("Only Admin is allowed", status=status.HTTP_400_BAD_REQUEST)  
            
##################################################################################################################################
class Convert_pdf_foodlist(APIView):


    def get(self, request): 

        #foodlist = FoodItem.objects.all().order_by('id')
        foodlist = food_utils.get_all_food_item_instances()
        table_data = {"foodlist":foodlist}
        html_string = render_to_string('pdf_foodlist.html', table_data)
        result = pdfkit.from_string(html_string)
        response = HttpResponse(result, content_type='application/pdf')
        response['Content-Disposition']='attachment;filename="foodlist.pdf"'
        return response
   

##################################################################################################################################

class Convert_pdf_by_category(APIView):


    def get(self, request): 

        category = request.query_params.get("category")
        if category=='Veg':
            foodlist = food_utils.filter_veg()
        elif category=='NonVeg':
            foodlist =food_utils.filter_nonveg()
        else: 
            return Response("This category is not available", status=status.HTTP_404_NOT_FOUND)   
        if foodlist.count() == 0:	
            return Response("No Food item in this category is available", status=status.HTTP_404_NOT_FOUND)   
        else: 
            table_data = {"foodlist":foodlist}
            html_string = render_to_string('pdf_foodlist.html', table_data)
            result = pdfkit.from_string(html_string)
            response = HttpResponse(result, content_type='application/pdf')
            response['Content-Disposition']='attachment;filename="foodlist.pdf"'
            return response
            

##################################################################################################################################

class Convert_pdf_by_calorieless(APIView):


    def get(self, request): 

        #foodlist = FoodItem.objects.all().filter(calories__lt=100) 
        foodlist = food_utils.calories_lt_100()
        if foodlist.count() == 0:	
            return Response("No Food item found", status=status.HTTP_404_NOT_FOUND)   
        else: 
            table_data = {"foodlist":foodlist}
            html_string = render_to_string('pdf_foodlist.html', table_data)
            result = pdfkit.from_string(html_string)
            response = HttpResponse(result, content_type='application/pdf')
            response['Content-Disposition']='attachment;filename="foodlist.pdf"'
            return response
            

##################################################################################################################################

        
class Convert_pdf_by_caloriegt(APIView):


    def get(self, request): 

        #foodlist = FoodItem.objects.all().filter(calories__gt=300)
        foodlist = food_utils.calories_gt_300()  
        if foodlist.count() == 0:	
            return Response("No Food item found", status=status.HTTP_404_NOT_FOUND)   
        else: 
            table_data = {"foodlist":foodlist}
            html_string = render_to_string('pdf_foodlist.html', table_data)
            result = pdfkit.from_string(html_string)
            response = HttpResponse(result, content_type='application/pdf')
            response['Content-Disposition']='attachment;filename="foodlist.pdf"'
            return response

##################################################################################################################################

class Convert_pdf_by_calorierange(APIView):


    def get(self, request): 

        #foodlist = FoodItem.objects.filter(calories__range=(100,200))  
        foodlist = food_utils.calories_range()
        
        if foodlist.count() == 0:	
            return Response("No Food item found", status=status.HTTP_404_NOT_FOUND)   
        else: 
            table_data = {"foodlist":foodlist}
            html_string = render_to_string('pdf_foodlist.html', table_data)
            result = pdfkit.from_string(html_string)
            response = HttpResponse(result, content_type='application/pdf')
            response['Content-Disposition']='attachment;filename="foodlist.pdf"'
            return response
            


##################################################################################################################################
#DOCUMENT GENERATION

class WordDocumentView(APIView):
    def get(self, request):
       
        document = Document() # Create a new document
        heading = document.add_heading('FOOD ITEMS', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        header_row = table.rows[0]
        header_row.cells[0].text = 'Sl. No.'
        header_row.cells[1].text = 'ID'
        header_row.cells[2].text = 'Item Name'
        header_row.cells[3].text = 'Category'
        header_row.cells[4].text = 'Quantity (grams)'
        header_row.cells[5].text = 'Calories (cals)'
       # Add some data to the table
        queryset = food_utils.get_all_food_item_instances()
        serializer = FoodSerializer(queryset, many=True)
        foodlist_list = serializer.data
        for index, row in enumerate(foodlist_list, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = str(row['id'])
            row_cells[2].text = row['name']
            row_cells[3].text = row['category']
            row_cells[4].text = str(row['quantity'])
            row_cells[5].text = str(row['calories'])
        print("document created successfully")

       # Save the document to a file
        document.save('djdocument.docx')
        with open('djdocument.docx', 'rb') as f:
            file_data = f.read()
        response = HttpResponse(file_data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="djdocument.docx"'
        return response


##################################################################################################################################

#TEMPLATE VIEWS

def foodlistpage(request):
    return render(request, 'foodlist.html')

def addfoodpage(request):
    return render(request, 'addfood.html')

def bulkcreatepage(request):
    return render(request, 'bulkcreate.html')

